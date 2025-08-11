import os
import json
import logging
import zipfile
from datetime import datetime
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from pydantic import BaseModel, Field
from typing import List
from langchain_core.runnables import RunnableParallel
import load_env
load_env()


api_key=os.env("OPENAI_API_KEY")
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load API key from .env

if not api_key:
    raise ValueError("OPENAI_API_KEY not found. Please set it in your .env file.")

project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

llm = ChatOpenAI(model="gpt-4o", temperature=0, openai_api_key=api_key)
embedding_model = OpenAIEmbeddings(openai_api_key=api_key)
vectorstore = FAISS.load_local(os.path.join(project_root, "adgm_vectorstore"), embedding_model, allow_dangerous_deserialization=True)
retriever = vectorstore.as_retriever()


class Issue(BaseModel):
    document: str = Field(description="The name of the document where the issue was found.")
    section: str = Field(description="The clause or section where the issue is located.")
    issue: str = Field(description="A detailed description of the red flag or inconsistency.")
    severity: str = Field(description="The severity of the issue (e.g., 'High', 'Medium', 'Low').")
    suggestion: str = Field(description="A suggestion for how to correct the issue, citing the relevant ADGM rule.")


class ValidationReport(BaseModel):
    issues_found: List[Issue] = Field(description="A list of all legal red flags and inconsistencies found.")
    summary: str = Field(description="A brief summary of the overall compliance status of the document(s).")


def parse_document(file_path):
    """Parses a .docx or .pdf file and returns its text content"""
    try:
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            content = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
        elif file_path.endswith('.pdf'):
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            content = ""
            for page in reader.pages:
                content += page.extract_text() or ""
        else:
            logger.error(f"Unsupported file type: {os.path.basename(file_path)}")
            return None
        
        logger.info(f"Successfully parsed {os.path.basename(file_path)}")
        return content
    except Exception as e:
        logger.error(f"Error parsing {file_path}: {e}")
        return None

# ... (rest of the functions remain the same) ...

def classify_document(document_content):
    """Enhanced document classification"""
    try:
        prompt_template = PromptTemplate(
            input_variables=["document_types", "document_content"],
            template="""You are a document classifier for ADGM legal documents. Your task is to identify the type of the following document from this predefined list: {document_types}.
            
            Document Content: {document_content}
            
            Please respond with only the document type and nothing else."""
        )
        
        document_types = [
            "Articles of Association", "Memorandum of Association", "Board Resolution", 
            "Incorporation Application Form", "UBO Declaration Form", "Employment Contract", 
            "Checklist - Private Company Limited", "Shareholder Resolution - Amendment of Articles", 
            "Licensing Regulatory Filings", "Commercial Agreements", "Compliance Risk Policies",
            "Financial Services Permission Application", "Banking License Application", "Insurance License Application",
            "Register of Members and Directors", "Change of Registered Address Notice"
        ]
        
        formatted_prompt = prompt_template.format(
            document_types=", ".join(document_types), 
            document_content=document_content[:2000]
        )
        response = llm.invoke(formatted_prompt)
        doc_type = response.content.strip()
        logger.info(f"Document classified as: {doc_type}")
        return doc_type
    except Exception as e:
        logger.error(f"Error classifying document: {e}")
        return "Unknown Document Type"


def rag_validate_document(file_path, document_content):
    """Enhanced RAG validation with better error handling"""
    try:
        output_parser = JsonOutputParser(pydantic_object=ValidationReport)
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert legal assistant specializing in ADGM regulations. Use the provided ADGM regulations as your sole source of truth to review the document for compliance. 

RED FLAGS TO DETECT:
- Invalid or missing clauses
- Incorrect jurisdiction (UAE Federal Courts instead of ADGM)
- Ambiguous or non-binding language
- Missing signatory sections or improper formatting
- Non-compliance with ADGM-specific templates

For each issue, cite the specific ADGM law or rule. Provide a brief summary of the document's compliance status. {format_instructions}"""),
            ("human", """ADGM Regulations: {context}\nDocument to Analyze: {document}\nPlease provide your analysis."""),
        ]).partial(format_instructions=output_parser.get_format_instructions())
        
        rag_chain = (
            RunnableParallel(
                context=lambda x: retriever.invoke(x['document']),
                document=lambda x: x['document']
            )
            | prompt
            | llm
            | output_parser
        )
        
        analysis = rag_chain.invoke({"document": document_content})
        
        for issue in analysis['issues_found']:
            issue['document'] = os.path.basename(file_path)
            
        logger.info(f"RAG validation completed: {len(analysis['issues_found'])} issues found")
        return analysis
        
    except Exception as e:
        logger.error(f"Error in RAG validation: {e}")
        return {"issues_found": [], "summary": f"Validation error: {str(e)}"}


def verify_checklist(classified_docs):
    """Enhanced checklist verification"""
    try:
        checklist_path = os.path.join(project_root, 'data', 'checklists.json')
        
        if not os.path.exists(checklist_path):
            os.makedirs(os.path.dirname(checklist_path), exist_ok=True)
            default_checklists = {
                "Company Incorporation": ["Articles of Association", "Memorandum of Association", "Incorporation Application Form", "UBO Declaration Form", "Register of Members and Directors"],
                "Financial Services Licensing": ["Financial Services Permission Application", "Business Plan", "Risk Management Framework", "Compliance Manual"],
                "HR & Employment": ["Employment Contract", "Employee Handbook", "Visa Application Form"],
                "Commercial Agreement Review": ["Commercial Agreements", "Board Resolution"],
                "Document Review": ["Any ADGM Document"]
            }
            with open(checklist_path, 'w') as f:
                json.dump(default_checklists, f, indent=2)
        
        with open(checklist_path, 'r') as f:
            checklists = json.load(f)

        process_prompt = PromptTemplate(
            input_variables=["classified_docs", "checklist_processes"],
            template="""You are a legal expert. Given the following classified documents, what legal process is the user attempting?
            
            Classified Documents: {classified_docs}
            Possible Processes: {checklist_processes}
            
            Respond with only the name of the process. If you cannot determine the process, respond with 'Unknown Process'."""
        )

        process_names = list(checklists.keys())
        try:
            process = llm.invoke(process_prompt.format(
                classified_docs=", ".join(classified_docs), 
                checklist_processes=", ".join(process_names)
            )).content.strip()
        except Exception as e:
            logger.error(f"Error in process detection: {e}")
            process = "Document Review"
        
        if process == "Unknown Process" or process not in checklists:
            return {"process": process, "documents_uploaded": len(classified_docs), "required_documents": "N/A", "missing_document": [], "issues_found": []}
            
        required_docs = checklists.get(process, [])
        uploaded_docs_set = set(classified_docs)
        missing_docs = []
        
        for required_doc in required_docs:
            found_match = False
            for uploaded_doc in uploaded_docs_set:
                if (uploaded_doc.lower() in required_doc.lower() or 
                    required_doc.lower() in uploaded_doc.lower() or
                    any(word in uploaded_doc.lower() for word in required_doc.lower().split()[:2] if len(word) > 3)):
                    found_match = True
                    break
            if not found_match:
                missing_docs.append(required_doc)
        
        result = {
            "process": process,
            "required_documents": len(required_docs),
            "documents_uploaded": len(uploaded_docs_set),
            "missing_document": missing_docs,
            "issues_found": []
        }
        
        logger.info(f"Checklist verification: {process} - {len(missing_docs)} missing documents")
        return result
        
    except Exception as e:
        logger.error(f"Error in checklist verification: {e}")
        return {"process": "Unknown Process", "documents_uploaded": len(classified_docs), "required_documents": 0, "missing_document": [], "issues_found": []}

def generate_output_docs(file_paths, validation_reports):
    """FIXED: Generate reviewed documents for ALL files and create zip archive"""
    try:
        output_docs_dir = os.path.join(project_root, "reviewed_docs")
        os.makedirs(output_docs_dir, exist_ok=True)
        
        final_json_report = {"status": "success", "processed_documents": []}
        reviewed_file_paths = []
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        for file_path, report in zip(file_paths, validation_reports):
            try:
                doc = Document(file_path)
                issues = report.get('issues_found', [])
                comments_added = 0
                
                logger.info(f"Processing {os.path.basename(file_path)} with {len(issues)} issues")
                
                for issue in issues:
                    found_comment_location = False
                    issue_section = issue.get('section', '')
                    
                    if not issue_section:
                        continue
                        
                    for paragraph in doc.paragraphs:
                        if not paragraph.text.strip():
                            continue
                            
                        if (issue_section.lower() in paragraph.text.lower() or
                            any(word in paragraph.text.lower() 
                                for word in issue_section.lower().split()[:3] if len(word) > 3)):
                            
                            if paragraph.runs:
                                severity = issue.get('severity', 'Low').lower()
                                if severity == 'high':
                                    paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.RED
                                elif severity == 'medium':
                                    paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                                else:
                                    paragraph.runs[0].font.highlight_color = WD_COLOR_INDEX.CYAN
                                
                                comment_text = f"""🚨 ADGM COMPLIANCE ISSUE

📋 Severity: {issue.get('severity', 'Unknown')}
📄 Section: {issue.get('section', 'Not specified')}
⚠️ Issue: {issue.get('issue', 'No description')}
💡 Suggestion: {issue.get('suggestion', 'No suggestion')}

Per ADGM Regulations - Generated by ADGM Corporate Agent"""
                                
                                comment = doc.add_comment(
                                    runs=paragraph.runs[0],
                                    text=comment_text,
                                    author='ADGM Corporate Agent',
                                    initials='ACA'
                                )
                                paragraph.runs[0].comment = comment
                                comments_added += 1
                                found_comment_location = True
                                break
                    
                    if not found_comment_location and doc.paragraphs and doc.paragraphs[0].runs:
                        fallback_comment = f"""🚨 ADGM COMPLIANCE ISSUE (General)

⚠️ Issue: {issue.get('issue', 'No description')}
💡 Suggestion: {issue.get('suggestion', 'No suggestion')}

Generated by ADGM Corporate Agent"""
                        
                        comment = doc.add_comment(
                            runs=doc.paragraphs[0].runs[0],
                            text=fallback_comment,
                            author='ADGM Corporate Agent',
                            initials='ACA'
                        )
                        doc.paragraphs[0].runs[0].comment = comment
                        comments_added += 1
                        
                original_file_name = os.path.basename(file_path)
                reviewed_file_name = f"reviewed_{timestamp}_{original_file_name}"
                reviewed_file_path = os.path.join(output_docs_dir, reviewed_file_name)
                doc.save(reviewed_file_path)
                
                reviewed_file_paths.append(reviewed_file_path)
                
                report_entry = {
                    "file_name": original_file_name,
                    "reviewed_file": reviewed_file_name,
                    "issues_count": len(issues),
                    "comments_added": comments_added,
                    "issues": issues
                }
                final_json_report["processed_documents"].append(report_entry)
                
                logger.info(f"✅ Generated reviewed document: {reviewed_file_name} with {comments_added} comments")
                
            except Exception as e:
                logger.error(f"❌ Error processing {file_path}: {e}")
                continue
        
        if reviewed_file_paths:
            zip_file_name = f"ADGM_Reviewed_Documents_{timestamp}.zip"
            zip_file_path = os.path.join(output_docs_dir, zip_file_name)
            
            with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in reviewed_file_paths:
                    zipf.write(file_path, os.path.basename(file_path))
                    
                json_report_name = f"ADGM_Analysis_Report_{timestamp}.json"
                json_report_path = os.path.join(output_docs_dir, json_report_name)
                with open(json_report_path, 'w') as f:
                    json.dump(final_json_report, f, indent=2)
                zipf.write(json_report_path, json_report_name)
            
            logger.info(f"📦 Created ZIP file: {zip_file_name} containing {len(reviewed_file_paths)} reviewed documents")
            
            final_json_report["zip_file"] = {
                "name": zip_file_name,
                "path": zip_file_path,
                "contains": len(reviewed_file_paths),
                "individual_files": [os.path.basename(f) for f in reviewed_file_paths]
            }
            
            return zip_file_path, final_json_report
        else:
            logger.warning("⚠️ No reviewed documents were generated")
            return None, {"status": "error", "message": "No documents could be processed successfully"}
        
    except Exception as e:
        logger.error(f"❌ Critical error in generate_output_docs: {e}")
        return None, {"status": "error", "message": str(e)}


def run_corporate_agent(file_paths):
    """Main ADGM Corporate Agent processing function - FIXED for multiple files"""
    try:
        logger.info(f"🚀 Starting ADGM processing for {len(file_paths)} files")
        
        parsed_docs = {}
        classified_docs = []
        
        for file_path in file_paths:
            content = parse_docx(file_path)
            if content:
                doc_type = classify_document(content)
                parsed_docs[file_path] = {
                    'content': content,
                    'type': doc_type
                }
                classified_docs.append(doc_type)
        
        if not parsed_docs:
            return None, {"status": "error", "message": "No documents could be processed"}
        
        logger.info(f"📄 Successfully processed {len(parsed_docs)} documents: {classified_docs}")
        
        checklist_result = verify_checklist(classified_docs)
        logger.info(f"📋 Process detected: {checklist_result.get('process', 'Unknown')}")
        
        validation_reports = []
        for file_path, doc_info in parsed_docs.items():
            validation = rag_validate_document(file_path, doc_info['content'])
            validation_reports.append(validation)
        
        total_issues = sum(len(report.get('issues_found', [])) for report in validation_reports)
        logger.info(f"🔍 Found {total_issues} total compliance issues across all documents")
        
        reviewed_output, file_report = generate_output_docs(
            list(parsed_docs.keys()),
            validation_reports
        )
        
        all_issues = []
        for report in validation_reports:
            all_issues.extend(report.get('issues_found', []))
        
        final_report = {
            "process": checklist_result.get('process', 'Unknown Process'),
            "documents_uploaded": checklist_result.get('documents_uploaded', len(classified_docs)),
            "required_documents": checklist_result.get('required_documents', 0),
            "missing_document": checklist_result.get('missing_document', []),
            "issues_found": all_issues,
            "total_documents_processed": len(parsed_docs),
            "documents_with_issues": len([r for r in validation_reports if r.get('issues_found')])
        }
        
        logger.info(f"✅ ADGM Corporate Agent processing completed successfully!")
        logger.info(f"📊 Summary: {len(parsed_docs)} docs processed, {len(all_issues)} issues found")
        
        return reviewed_output, final_report
        
    except Exception as e:
        logger.error(f"❌ Critical error in ADGM processing: {e}")
        return None, {"status": "error", "message": str(e)}